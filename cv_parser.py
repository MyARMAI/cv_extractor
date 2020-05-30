from docx import Document
import pandas as pd

test_file = r"C:\Users\Cheikh\Desktop\Projet_memoire\myArmAi\samples\cv\cv_atos\fr\word\CV ATOS SIMON PIERRE DIOUF.docx"

document = Document(test_file)
table = document.tables[0]

section = [""]


# transform word table to pandas dataframe
def tableToDF(wordDoc):
    tables = []
    for table in wordDoc.tables:
        df = [['' for i in range(len(table.columns))]
              for j in range(len(table.rows))]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    df[i][j] = cell.text
        tables.append(pd.DataFrame(df))
    return tables


if __name__ == "__main__":
    res = tableToDF(document)
    len(res)
