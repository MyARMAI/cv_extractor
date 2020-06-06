#!/usr/bin/env python
# coding: utf-8

import string
import win32com.client as win32
import textract
import os
import pandas as pd
import re


from collections import Counter
from docx import Document
from win32com.client import constants
from glob import glob

import pythoncom


pd.set_option('display.max_rows', 500)
pd.set_option('display.max_columns', 500)
pd.set_option('display.width', 1000)
pd.set_option('display.max_colwidth', None)


def renameDoc2Docx(path):
    dir = os.listdir(path)
    for d in dir:
        file = os.path.join(path, d)
        if(os.path.isfile(file)):
            name, extension = file.rsplit(".")
            if(extension == "doc"):
                new_name = name + '.docx'
                os.rename(file, new_name)


def save_as_docx(path):
    # Opening MS Word
    # needed because of flask multithreading politics
    pythoncom.CoInitialize()
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)
    return new_file_abs


def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell


def readFile(filepath):
    data = []
    filename, extension = filepath.rsplit(".")
    _file = filepath
    # in windows system not checking this may cause error if one the files have been opened
    if not re.search(r'\$', filename):
        if extension == "doc":
            path = glob(r'..\temp\*.doc')
            _file = save_as_docx(filepath)

        document = Document(_file)
        table = document.tables[0]
        res = tableToDF(document)
        data.append(res)
    return data

# transform word table to pandas dataframe


def tableToDF(wordDoc):
    tables = []
    for table in wordDoc.tables:
        df = [['' for i in range(len(table.columns))]
              for j in range(len(table.rows))]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(iter_unique_cells(row)):
                if cell.text:
                    df[i][j] = cell.text
        tables.append(pd.DataFrame(df))
    return tables


def dropSpace(s, replace=""):
    return re.sub(r'\n', replace, s)


printable = set(string.printable)


# get cv language
def getLang(df):
    patterns = ["language", "langue"]
    dfSize = len(df)
    lang = ""
    for i in range(dfSize):
        _df = df[i]
        for index, row in _df.iterrows():
            if(re.search("language", row[0].lower()) or re.search("langue", row[0].lower())):
                if(len(row) == 3):
                    lang = row[2]
                elif len(row) == 2:
                    lang = row[1]
    lang = ''.join(filter(lambda x: x in printable, lang))
    lang = dropSpace(lang, replace=' ')
    return lang


def getNameAndProfil(df):
    data = df[0]
    name = ''
    profil = ''
    for i in data:
        size = len(data[i])
        if size > 0 and size < 5:
            try:
                if(int(data[i].str.index("("))):  # contain parenthesis xxx (X olds) XXX

                    profil_ind = int(data[i].str.index(")"))
                    name_ind = int(data[i].str.index("("))
                    val = data[i][0]

                    name = val[0:name_ind-1]
                    profil = val[profil_ind+1:]

                    ## val = re.sub("[^\w]"," ",val)
                    name = re.sub("[\\r\\n\|\\t|[0-9]]*", "", name)
                    ##profil = re.sub("[\\r\\n\|\\t|[0-9]]*", "", profil)
                    profil = [x for x in profil.split("\n") if len(x) > 0]
                    # print(len(profil))

                    if len(profil) == 2:

                        # search a string that doesn't contain any number followed by space
                        if re.search(r'[^0-9]\s', profil[1]):
                            profil = profil[0]
                        else:
                            profil = profil[0]
                           # print(profil)
                    else:
                        profil = profil[0]
                       # print(profil)

            except:
                value = data[i].tolist()[0].split("\n\n")
                if len(value) == 2:
                    name, profil = value
                    name = dropSpace(name)
                    profil = re.sub("\t", "", profil)
                    # print(profil)
                elif len(value[0]) != 0:  # some cv may contain nom prenom and profil
                    val = value[0].split(":")
                    if len(val) > 2:
                        name, profil = val[1].split("\n")[0], val[2]
                        # print(name,profil)
                pass

    name = name.strip()
    profil = profil.strip()
    return name, profil


def getFormation(data):
    pattern = ["formation", "training"]
    dfSize = len(data)
    formation = ""

    for i in range(dfSize):
        _data = data[i]

        for index, row in _data.iterrows():

            if(re.search(pattern[0], row[0].lower()) or re.search(pattern[1], row[0].lower())):
                if(len(row) == 3):
                    formation = row[2]

    formation = dropSpace(formation, " ")
    formation = re.sub(' +', " ", formation)
    return formation


def getExperience(data):
    pattern = ["experience", "expérience"]
    dfSize = len(data)
    experience = ""
    for i in range(dfSize):
        _data = data[i]
        for index, row in _data.iterrows():
            if len(row) <= 3:
                for el in row:
                    if re.search(pattern[0], el.lower()) or re.search(pattern[1], el.lower()):
                        el = re.sub(r'\((.*?)\)', '', el)
                        res = re.search(r'[0-9]{1,2}\s', el)
                        if res != None:
                            begin, end = res.span()
                            experience = el[begin:end]
                            # print(experience)
    return experience


def getSkills(data):
    pattern = ["compétences", "skills"]
    dfSize = len(data)
    competences = {}
    _max = max(Counter([len(x) for x in data]).keys())

    for i in range(dfSize):
        _data = data[i]
        for index, row in _data.iterrows():
            # find principales and fonctionnale competences
            for el in row:
                if re.search(pattern[0], el.lower()) or re.search(pattern[1], el.lower()):
                    value = [x for x in row.tolist() if len(x) != 0]
                    out = ""
                    for s in value:
                        if(not re.search(pattern[0], s.lower())):
                            out += s+u" "
                    competences[el] = out

    competences = {x: dropSpace(y.replace("\xa0", ""), replace=' ')
                   for x, y in competences.items()}
    # print(competences)
    # find technical competences
    # we assume that array containing technical skills have the same length
    # print(row)
    skills = [x for x in data if len(x) == _max][0]
    return skills, competences


def parser(filepath):
    data = readFile(filepath)[0]
    result = {}
    result["name"], result["profil"] = getNameAndProfil(data)
    result["lang"] = getLang(data)
    result["experience"] = getExperience(data)
    result["formation"] = getFormation(data)
    skills, result["competences"] = getSkills(data)
    skills = skills.to_dict()
    #print("Name - {} Profil - {} Expérience - {}ans ".format(name, profil, exp))
    # print('Language')
    # print(lang)
    # print("Formation")
    #print(re.sub(' +', " ", formation))
    _skills = []
    for key, _dict in skills.items():
        for k, v in _dict.items():
            if(len(v) != 0):
                _skills.append(re.sub(' +', " ", v))
    result["skills"] = _skills
    #print(type(data), type(name), type(profil), type(lang), type(exp), type(formation), type(_skills))
    return result


if __name__ == "__main__":
    test_file = r'C:\Users\Cheikh\Desktop\Projet_memoire\myArmAi\samples\base_cv\cv\CV ATOS Amadou NDIAYE - ENGLISH.docx'
    parser(test_file)
